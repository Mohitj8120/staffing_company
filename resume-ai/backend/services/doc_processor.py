import os
from docx import Document
from docxtpl import DocxTemplate
from docx2pdf import convert
import fitz  # PyMuPDF
from jinja2 import Template
from playwright.sync_api import sync_playwright
from core.config import settings
import subprocess
import copy

def sanitize_contact_links(data: dict) -> dict:
    """
    Ensures that linkedin, github, and portfolio links in data['personal'] 
    start with http:// or https://. Returns a copy of the data dictionary.
    """
    data_copy = copy.deepcopy(data)
    if "personal" in data_copy and isinstance(data_copy["personal"], dict):
        for field in ["linkedin", "github", "portfolio"]:
            val = data_copy["personal"].get(field)
            if val and isinstance(val, str):
                val_stripped = val.strip()
                if val_stripped and not (val_stripped.startswith("http://") or val_stripped.startswith("https://")):
                    data_copy["personal"][field] = "https://" + val_stripped
    return data_copy

def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts all text from a DOCX file.
    """
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also extract text from tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    return '\n'.join(full_text)

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts all text from a PDF file using PyMuPDF.
    """
    doc = fitz.open(file_path)
    full_text = []
    for page in doc:
        full_text.append(page.get_text("text"))
        # Also extract links so LLM can parse LinkedIn/GitHub URLs
        for link in page.get_links():
            if 'uri' in link:
                full_text.append(link['uri'])
    return '\n'.join(full_text)

def start_pdf_render_process(output_filename: str) -> subprocess.Popen:
    """
    Asynchronously starts the Playwright rendering subprocess in the background.
    Imports Playwright and launches Chromium early, then blocks waiting for HTML on stdin.
    """
    import subprocess
    import sys
    output_path = os.path.join(settings.TEMP_DIR, output_filename)
    
    script = """
import sys
from playwright.sync_api import sync_playwright
sys.stdin.reconfigure(encoding='utf-8')
with sync_playwright() as p:
    browser = p.chromium.launch(headless=True)
    page = browser.new_page()
    page.set_content(sys.stdin.read())
    page.pdf(path=sys.argv[1], format="A4", margin={"top": "20px", "right": "20px", "bottom": "20px", "left": "20px"}, print_background=True)
    browser.close()
"""
    
    process = subprocess.Popen(
        [sys.executable, "-c", script, output_path],
        stdin=subprocess.PIPE,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        encoding="utf-8"
    )
    return process

def trim_resume_data(data: dict, target_pages: int, severity: int = 1):
    """
    Trims resume data to help it fit into target pages.
    severity 1: mild pruning (limit bullet points slightly, remove excess items).
    severity 2: aggressive pruning (strict limits on items, fewer bullet points).
    """
    if target_pages == 1:
        # Experience
        if "experience" in data and isinstance(data["experience"], list):
            # Limit number of experience roles
            max_roles = 3 if severity > 1 else 4
            data["experience"] = data["experience"][:max_roles]
            for exp in data["experience"]:
                if "points" in exp and isinstance(exp["points"], list):
                    max_points = 3 if severity > 1 else 4
                    exp["points"] = exp["points"][:max_points]
        # Projects
        if "projects" in data and isinstance(data["projects"], list):
            max_projs = 2 if severity > 1 else 3
            data["projects"] = data["projects"][:max_projs]
            for proj in data["projects"]:
                if "points" in proj and isinstance(proj["points"], list):
                    max_points = 2 if severity > 1 else 3
                    proj["points"] = proj["points"][:max_points]
        # Certifications
        if "certifications" in data and isinstance(data["certifications"], list):
            max_certs = 3 if severity > 1 else 4
            data["certifications"] = data["certifications"][:max_certs]
            
    elif target_pages == 2:
        # Experience
        if "experience" in data and isinstance(data["experience"], list):
            max_roles = 4 if severity > 1 else 5
            data["experience"] = data["experience"][:max_roles]
            for exp in data["experience"]:
                if "points" in exp and isinstance(exp["points"], list):
                    max_points = 4 if severity > 1 else 5
                    exp["points"] = exp["points"][:max_points]
        # Projects
        if "projects" in data and isinstance(data["projects"], list):
            max_projs = 3 if severity > 1 else 4
            data["projects"] = data["projects"][:max_projs]
            for proj in data["projects"]:
                if "points" in proj and isinstance(proj["points"], list):
                    max_points = 3 if severity > 1 else 4
                    proj["points"] = proj["points"][:max_points]
        # Certifications
        if "certifications" in data and isinstance(data["certifications"], list):
            max_certs = 4 if severity > 1 else 6
            data["certifications"] = data["certifications"][:max_certs]

def generate_stunning_pdf(
    data: dict, 
    output_filename: str = "tailored_resume.pdf", 
    mode: str = "standard", 
    page_count: str = "auto",
    prewarmed_process: subprocess.Popen = None
) -> str:
    """
    Renders the HTML template and writes it to Playwright (either direct launch or pre-warmed),
    with an adaptive loop checking the output PDF page count and compacting/trimming to fit constraints.
    """
    data = sanitize_contact_links(data)
    # Determine the target page limit
    target_page_count = 2 # Default fallback
    if page_count == "1":
        target_page_count = 1
    elif page_count == "2":
        target_page_count = 2
    elif page_count == "auto":
        orig_pc = data.get("original_page_count", 2)
        try:
            target_page_count = min(int(orig_pc), 2)
        except:
            target_page_count = 2

    # If target is 1 page, apply an initial mild trim
    if target_page_count == 1:
        trim_resume_data(data, 1, severity=1)

    compactor_levels = ["", "compact-level-1", "compact-level-2", "compact-level-3"]
    compactor_levels = ["", "compact-level-1", "compact-level-2", "compact-level-3"]
    current_prewarmed = prewarmed_process
    output_path = os.path.join(settings.TEMP_DIR, output_filename)
    success = False
    
    # Start inline browser if no prewarmed process is provided to avoid subprocess overhead
    browser = None
    playwright_context = None
    inline_page = None
    
    if current_prewarmed is None:
        try:
            from playwright.sync_api import sync_playwright
            playwright_context = sync_playwright().start()
            browser = playwright_context.chromium.launch(headless=True)
            inline_page = browser.new_page()
            print("PDF Page Fitting: Successfully pre-warmed inline browser for fitting loops.")
        except Exception as e:
            print(f"Warning: Failed to launch inline Playwright: {e}. Will fallback to slower subprocess method.")
            
    # Try compaction and trimming iteratively
    for severity in [0, 1, 2]:
        if severity > 0:
            print(f"PDF Page Fitting: Trimming data at severity {severity} to fit {target_page_count} page(s)...")
            trim_resume_data(data, target_page_count, severity=severity)
            
        for level in compactor_levels:
            # Render HTML template with current compactor level
            template_name = "resume_redesign_theme.html" if mode == "redesign" else "resume_theme.html"
            template_path = os.path.join(settings.BASE_DIR, "templates", template_name)
            if not os.path.exists(template_path):
                template_path = os.path.join(os.path.dirname(__file__), "..", "templates", template_name)
                
            with open(template_path, "r", encoding="utf-8") as f:
                template_str = f.read()
                
            template = Template(template_str)
            html_content = template.render(data=data, page_count=page_count, compactor_class=level)
            
            # Minify HTML to save memory and optimize processing speed
            import re
            def minify_html(html: str) -> str:
                html = re.sub(r'<!--.*?-->', '', html, flags=re.DOTALL)
                html = re.sub(r'[\r\n\t]+', ' ', html)
                html = re.sub(r'\s{2,}', ' ', html)
                return html.strip()
            
            html_content = minify_html(html_content)
            
            import subprocess
            import sys
            
            # Render PDF using current Playwright process
            if current_prewarmed is not None:
                try:
                    stdout, stderr = current_prewarmed.communicate(input=html_content)
                    if current_prewarmed.returncode != 0:
                        print(f"Warning: Playwright pre-warmed PDF generation failed: {stderr}")
                        current_prewarmed = None
                except Exception as e:
                    print(f"Warning: Exception using pre-warmed process: {e}")
                    current_prewarmed = None
                    
            if current_prewarmed is None:
                if inline_page is not None:
                    try:
                        inline_page.set_content(html_content)
                        inline_page.pdf(path=output_path, format="A4", margin={"top": "20px", "right": "20px", "bottom": "20px", "left": "20px"}, print_background=True)
                    except Exception as inline_err:
                        print(f"Inline PDF render failed: {inline_err}. Falling back to subprocess...")
                        script = """
import sys
from playwright.sync_api import sync_playwright
sys.stdin.reconfigure(encoding='utf-8')
with sync_playwright() as p:
    browser = p.chromium.launch(headless=True)
    page = browser.new_page()
    page.set_content(sys.stdin.read())
    page.pdf(path=sys.argv[1], format="A4", margin={"top": "20px", "right": "20px", "bottom": "20px", "left": "20px"}, print_background=True)
    browser.close()
"""
                        process = subprocess.run(
                            [sys.executable, "-c", script, output_path],
                            input=html_content,
                            text=True,
                            encoding="utf-8",
                            capture_output=True
                        )
                        if process.returncode != 0:
                            raise Exception(f"Playwright PDF generation failed: {process.stderr}")
                else:
                    # Fallback to subprocess directly
                    script = """
import sys
from playwright.sync_api import sync_playwright
sys.stdin.reconfigure(encoding='utf-8')
with sync_playwright() as p:
    browser = p.chromium.launch(headless=True)
    page = browser.new_page()
    page.set_content(sys.stdin.read())
    page.pdf(path=sys.argv[1], format="A4", margin={"top": "20px", "right": "20px", "bottom": "20px", "left": "20px"}, print_background=True)
    browser.close()
"""
                    process = subprocess.run(
                        [sys.executable, "-c", script, output_path],
                        input=html_content,
                        text=True,
                        encoding="utf-8",
                        capture_output=True
                    )
                    if process.returncode != 0:
                        raise Exception(f"Playwright PDF generation failed: {process.stderr}")
            
            # Check how many pages the generated PDF actually has!
            try:
                doc = fitz.open(output_path)
                actual_page_count = len(doc)
                doc.close()
                print(f"PDF Page Fitting check: level='{level}', severity={severity} -> rendered {actual_page_count} page(s) (target: {target_page_count})")
                
                if actual_page_count <= target_page_count:
                    success = True
                    break
            except Exception as fitz_err:
                print(f"Error checking PDF page count: {fitz_err}")
                success = True
                break
                
            # If we used prewarmed_process, it is now consumed. Clear it so next loop iteration uses direct launch.
            current_prewarmed = None
            
        if success:
            break
            
    # Cleanup pre-warmed inline browser resources
    if browser is not None:
        try:
            browser.close()
        except:
            pass
    if playwright_context is not None:
        try:
            playwright_context.stop()
        except:
            pass
            
    return output_path

def generate_tailored_docx(data: dict, output_filename: str = "tailored_resume.docx") -> str:
    """
    Fills the resume_template.docx with tailored data and saves it.
    Returns the path to the saved DOCX file.
    """
    data = sanitize_contact_links(data)
    template_path = os.path.join(settings.TEMPLATES_DIR, "resume_template.docx")
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at {template_path}")
        
    doc = DocxTemplate(template_path)
    
    # Format data for the template
    # The template expects specific placeholders like {{name}}, {{summary}}, {{experience}}
    
    experience_list = []
    for e in data.get("experience", []):
        if not isinstance(e, dict):
            continue
        title = e.get("title", "")
        company = e.get("company", "")
        duration = e.get("duration", "")
        points = e.get("points", [])
        points_str = "\n".join(["• " + str(p) for p in points]) if isinstance(points, list) else ""
        exp_str = f'{title} | {company} | {duration}\n' + points_str
        experience_list.append(exp_str)
    experience_text = "\n\n".join(experience_list)
    
    projects_list = []
    for p in data.get("projects", []):
        if not isinstance(p, dict):
            continue
        title = p.get("title", "")
        desc = p.get("description", p.get("summary", ""))
        if isinstance(desc, list):
            desc = "\n".join(["• " + str(x) for x in desc])
        projects_list.append(f'{title}\n{desc}')
    projects_text = "\n\n".join(projects_list)
    
    education_list = []
    for e in data.get("education", []):
        if not isinstance(e, dict):
            continue
        degree = e.get("degree", "")
        college = e.get("college", e.get("school", ""))
        year = e.get("year", "")
        education_list.append(f'{degree} | {college} | {year}')
    education_text = "\n\n".join(education_list)

    # Safely handle skills and certifications which might be string or list
    skills = data.get("skills", [])
    if isinstance(skills, list):
        skills_text = "\n".join([str(s) for s in skills])
    else:
        skills_text = str(skills)

    certs = data.get("certifications", [])
    if isinstance(certs, list):
        certs_text = "\n".join([str(c) for c in certs])
    else:
        certs_text = str(certs)

    context = {
        "name": data.get("personal", {}).get("name", "") if isinstance(data.get("personal"), dict) else "",
        "email": data.get("personal", {}).get("email", "") if isinstance(data.get("personal"), dict) else "",
        "phone": data.get("personal", {}).get("phone", "") if isinstance(data.get("personal"), dict) else "",
        "location": data.get("personal", {}).get("location", "") if isinstance(data.get("personal"), dict) else "",
        "linkedin": data.get("personal", {}).get("linkedin", "") if isinstance(data.get("personal"), dict) else "",
        "portfolio": data.get("personal", {}).get("portfolio", "") if isinstance(data.get("personal"), dict) else "",
        "summary": data.get("summary", ""),
        "skills": skills_text,
        "experience": experience_text,
        "projects": projects_text,
        "education": education_text,
        "certifications": certs_text
    }
    
    doc.render(context)
    
    output_path = os.path.join(settings.TEMP_DIR, output_filename)
    doc.save(output_path)
    
    return output_path

def convert_docx_to_pdf(docx_path: str, output_filename: str = "tailored_resume.pdf") -> str:
    """
    Converts a DOCX file to PDF using docx2pdf.
    Note: Requires Microsoft Word to be installed on Windows.
    """
    pdf_path = os.path.join(settings.TEMP_DIR, output_filename)
    convert(docx_path, pdf_path)
    return pdf_path

def get_docx_text_map(file_path: str):
    """
    Iterates through a DOCX and creates a map of paragraph indexes to their text.
    Returns (elements list, dict mapping str(idx) -> text).
    """
    doc = Document(file_path)
    elements = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            elements.append(para)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        elements.append(para)
                        
    text_map = {str(i): el.text for i, el in enumerate(elements)}
    return elements, text_map, doc

def apply_text_map_to_docx(file_path: str, modifications: dict, output_filename: str) -> str:
    """
    Applies text modifications to a DOCX file while retaining paragraph/first-run formatting.
    """
    elements, _, doc = get_docx_text_map(file_path)
    
    for idx_str, new_text in modifications.items():
        try:
            idx = int(idx_str)
            if idx < len(elements):
                para = elements[idx]
                # Split new text by newlines just in case the LLM outputs them, to avoid breaking Word layout
                parts = new_text.split('\n')
                
                if not para.runs:
                    for i, part in enumerate(parts):
                        para.add_run(part)
                        if i < len(parts) - 1:
                            para.add_run().add_break()
                else:
                    # Keep the formatting of the first run
                    first_run = para.runs[0]
                    first_run.text = parts[0]
                    
                    # If there are more parts due to newlines, append them to the first run's paragraph
                    for i in range(1, len(parts)):
                        first_run.add_break()
                        first_run.add_text(parts[i])
                        
                    # Clear out the rest of the runs
                    for i in range(1, len(para.runs)):
                        para.runs[i].text = ""
        except (ValueError, KeyError):
            pass
            
    # Apply "keep with next" for short headers to prevent page break issues
    for para in doc.paragraphs:
        # Keep lines together for all paragraphs so they don't break across pages
        para.paragraph_format.keep_lines_together = True
        
        # If it's a short heading or title, keep it with the next paragraph
        text = para.text.strip()
        if text and len(text.split()) <= 10:
            if para.style.name.startswith("Heading") or "Title" in para.style.name:
                para.paragraph_format.keep_with_next = True
            elif "Tech Stack" in text:
                para.paragraph_format.keep_with_next = True
            
    output_path = os.path.join(settings.TEMP_DIR, output_filename)
    doc.save(output_path)
    return output_path
